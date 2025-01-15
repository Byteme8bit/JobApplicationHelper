import unittest
import os
from utils import build_config_from_template, extract_placeholders
from docx import Document
import tempfile
import shutil


class TestUtils(unittest.TestCase):

    def setUp(self):
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.test_dir)

    def test_build_config_from_template_docx(self):
        # Create a temporary DOCX file for testing
        template_path = os.path.join(self.test_dir, "test-template.docx")
        doc = Document()
        doc.add_paragraph("This is a %placeholder% template.")
        doc.save(template_path)

        config_path = build_config_from_template(template_path)
        self.assertTrue(os.path.exists(config_path))
        with open(config_path, 'r') as f:
            config_data = json.load(f)
        self.assertEqual(config_data["templateFilePath"], template_path)
        self.assertIn("placeholder", config_data["placeholders"])

    def test_build_config_from_template_txt(self):
        # Create a temporary TXT file for testing
        template_path = os.path.join(self.test_dir, "test-template.txt")
        with open(template_path, "w") as f:
            f.write("This is a %placeholder% template.")

        config_path = build_config_from_template(template_path)
        self.assertTrue(os.path.exists(config_path))
        with open(config_path, 'r') as f:
            config_data = json.load(f)
        self.assertEqual(config_data["templateFilePath"], template_path)
        self.assertIn("placeholder", config_data["placeholders"])

    def test_extract_placeholders_docx(self):
        # Create a temporary DOCX file for testing
        template_path = os.path.join(self.test_dir, "test-template.docx")
        doc = Document()
        doc.add_paragraph("This is a %placeholder1% and %placeholder2% template.")
        doc.save(template_path)

        placeholders = extract_placeholders(template_path, "%")
        self.assertEqual(len(placeholders), 2)
        self.assertIn("placeholder1", placeholders)
        self.assertIn("placeholder2", placeholders)

    def test_extract_placeholders_txt(self):
        # Create a temporary TXT file for testing
        template_path = os.path.join(self.test_dir, "test-template.txt")
        with open(template_path, "w") as f:
            f.write("This is a %placeholder1% and %placeholder2% template.")

        placeholders = extract_placeholders(template_path, "%")
        self.assertEqual(len(placeholders), 2)
        self.assertIn("placeholder1", placeholders)
        self.assertIn("placeholder2", placeholders)


if __name__ == '__main__':
    unittest.main()

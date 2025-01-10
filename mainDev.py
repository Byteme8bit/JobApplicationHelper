import argparse
import json
import os
from datetime import date
from docx import Document
from gui import load_config_file, browse_config_file, browse_template_file, add_placeholder

def generate_document(template_path, output_path, data):
    """Generates a document, replacing placeholders and resizing text to fit.

    Supports .docx and .txt files.

    Args:
        template_path: Path to the template file.
        output_path: Path to save the generated document.
         Dictionary containing placeholder names and their replacements.
    """
    try:
        if template_path.endswith(".docx"):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    paragraph.text = paragraph.text.replace(f"%{key}%", value)
            doc.save(output_path)

        elif template_path.endswith(".txt"):
            with open(template_path, "r") as infile, open(output_path, "w") as outfile:
                for line in infile:
                    for key, value in data.items():
                        line = line.replace(f"%{key}%", value)
                    outfile.write(line)
    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {template_path}")
    except ValueError as e:
        raise
    except Exception as e:
        raise IOError(f"Error processing template: {e}")

    # Check for unmatched placeholders
    if template_path.endswith(".docx"):
        unmatched = [key for paragraph in doc.paragraphs for key in data if f"%{key}%" in paragraph.text]
    elif template_path.endswith(".txt"):
        with open(output_path, "r") as outfile:
            unmatched = [key for line in outfile for key in data if f"%{key}%" in line]
    if unmatched:
        raise ValueError(f"Unmatched placeholders found: {', '.join(unmatched)}")


def load_config(path):
    """Loads config from a JSON file. Returns an empty dictionary if the file doesn't exist."""
    try:
        with open(path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


if __name__ == "__main__":
    # Parse command-line arguments
    parser = argparse.ArgumentParser(description="Generate documents from templates.")
    parser.add_argument("--config", help="Path to the config file (default: config.json)")
    parser.add_argument("-GUI", action="store_true", help="Load the GUI")
    args = parser.parse_args()

    if args.GUI:
        from run_gui import run_gui
        run_gui() # Call the GUI function from the imported module
    else:
        config_path = args.config if args.config else input(
            "Enter the path to the config file (default: config.json): ") or "config.json"
        print(f"Loading config from {config_path}...")
        config = load_config(config_path)

        # Automatically generate date
        today = date.today().strftime("%Y-%m-%d")

        while True:
            template_path = config.get("templateFilePath")
            if template_path is None:
                template_path = input("Enter template file path: ")
            if not os.path.exists(template_path):
                print("Error: Template file not found. Please enter a valid path.")
                if template_path is not None:  # Only clear the config value if it was from the file
                    config["template_path"] = None  # Clear incorrect value from config
                continue  # Ask for input again
            print(f"Using template file: {template_path}")
            break

        while True:
            output_filename = config.get("outputFilePath")
            if output_filename is None:
                output_filename = input("Enter output file name: ")

            overwrite = config.get("overwriteOutput")  # Initialize overwrite variable
            if overwrite is None:
                overwrite = False

            while os.path.exists(output_filename):
                overwrite = input(f"File '{output_filename}' already exists. Overwrite? (y/n): ")
                if overwrite:
                    break  # Exit the loop if user wants to overwrite
                elif not overwrite:
                    output_filename = input("Enter a different output file name: ")  # Ask for a new name
                    overwrite = None  # Reset overwrite for the new filename check
                else:
                    print("Invalid input. Please enter 'y' or 'n'.")  # Handle invalid input

            if overwrite and output_filename is not None:
                break  # Exit the loop if user wants to overwrite
            else:
                continue  # Ask for input again if not overwriting

        print(f"Output file will be saved as: {output_filename}")

        placeholders = config.get("placeholders", {})  # Load placeholders from config, default to empty dict

        if not placeholders:  # Only ask for placeholders if data is empty
            print("No data found in config. Please enter placeholder values.")
            while True:
                key = input("Enter placeholder name (or type 'done'): ")
                if key == "done":
                    break
                value = input(f"Enter value for {key}: ")
                placeholders[key] = value

        # Add the automatically generated date to the data dictionary
        placeholders["date"] = today

        try:
            print("Generating document...")
            generate_document(template_path, output_filename, placeholders)
            print(f"Document '{output_filename}' generated successfully.")
        except Exception as e:
            print(f"Error: {e}")
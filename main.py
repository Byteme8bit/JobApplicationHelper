import argparse
import json
import os
from datetime import date
from docx import Document


def generate_document(template_path, output_path, data):
    """Generates a document, replacing placeholders and resizing text to fit.

    Supports .docx and .txt files.

    Args:
        template_path: Path to the template file.
        output_path: Path to save the generated document.
        data: Dictionary containing placeholder names and their replacements.
    """
    try:
        if template_path.endswith(".docx"):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    paragraph.text = paragraph.text.replace(f"%{key}%", value)
            doc.save(output_path)

        elif template_path.endswith(".txt"):
            with open(template_path, "r") as infile, open(output_path, "w") as outfile:
                for line in infile:
                    for key, value in data.items():
                        line = line.replace(f"%{key}%", value)
                    outfile.write(line)
    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {template_path}")
    except ValueError as e:
        raise
    except Exception as e:
        raise IOError(f"Error processing template: {e}")

    # Check for unmatched placeholders
    if template_path.endswith(".docx"):
        unmatched = [key for paragraph in doc.paragraphs for key in data if f"%{key}%" in paragraph.text]
    elif template_path.endswith(".txt"):
        with open(output_path, "r") as outfile:
            unmatched = [key for line in outfile for key in data if f"%{key}%" in line]
    if unmatched:
        raise ValueError(f"Unmatched placeholders found: {', '.join(unmatched)}")


def load_config(path):
    """Loads config from a JSON file. Returns an empty dictionary if the file doesn't exist."""
    try:
        with open(path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate documents from templates.")
    parser.add_argument("--config", help="Path to the config file (default: config.json)")
    args = parser.parse_args()

    config_path = args.config if args.config else input(
        "Enter the path to the config file (default: config.json): ") or "config.json"
    config = load_config(config_path)

    # Automatically generate date
    today = date.today().strftime("%Y-%m-%d")

    while True:
        template_path = config.get("templateFilePath")
        if template_path is None:
            template_path = input("Enter template file path: ")
        if not os.path.exists(template_path):
            print("Error: Template file not found. Please enter a valid path.")
            if template_path is not None:  # Only clear the config value if it was from the file
                config["template_path"] = None  # Clear incorrect value from config
            continue  # Ask for input again
        break

    while True:
        output_filename = config.get("outputFilePath")
        if output_filename is None:
            output_filename = input("Enter output file name: ")

        overwrite = config.get("overwriteOutput") if not None else False  # Initialize overwrite variable

        while os.path.exists(output_filename):
            overwrite = input(f"File '{output_filename}' already exists. Overwrite? (y/n): ")
            if overwrite.lower() == 'y' or overwrite:
                break  # Exit the loop if user wants to overwrite
            elif overwrite.lower() == 'n' or not overwrite:
                output_filename = input("Enter a different output file name: ")  # Ask for a new name
                overwrite = None  # Reset overwrite for the new filename check
            else:
                print("Invalid input. Please enter 'y' or 'n'.")  # Handle invalid input

        if (overwrite.lower() == 'y' or overwrite) and output_filename is not None:
            break  # Exit the loop if user wants to overwrite
        else:
            continue  # Ask for input again if not overwriting

    data = config.get("data", {})  # Load data from config, default to empty dict

    if not data:  # Only ask for placeholders if data is empty
        while True:
            key = input("Enter placeholder name (or type 'done'): ")
            if key == "done":
                break
            value = input(f"Enter value for {key}: ")
            data[key] = value

    # Add the automatically generated date to the data dictionary
    data["date"] = today

    try:
        generate_document(template_path, output_filename, data)
        print(f"Document '{output_filename}' generated successfully.")
    except Exception as e:
        print(f"Error: {e}")

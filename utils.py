import json
import os
import re
from docx import Document
from datetime import date


def open_config(config_filepath):
    """Opens the config file in default application native to OS"""
    try:
        # Use os.startfile for better cross-platform compatibility
        os.startfile(config_filepath)
        input("Press Enter after reviewing the config file...")  # Pause execution until user is ready
        return load_config(config_filepath)

    except FileNotFoundError:
        print(f"Error: Config file '{config_filepath}' not found.")
        return None  # Return None to indicate failure
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None  # Return None to indicate failure


def generate_document(config, bookend="%"):
    """Generates a document from a template and placeholders."""
    if not isinstance(config, dict):
        raise TypeError("Config must be a dictionary.")

    template_path = config.get("templateFilePath")
    output_path = config.get("outputFilePath")
    placeholders = config.get("placeholders", {})  # Handle missing placeholders gracefully

    try:
        if os.path.exists(output_path):
            if not config.get("overwriteOutput", False):
                overwrite = input(
                    f"Output file '{output_path}' already exists. Overwrite? (y/n): ").lower()
                if overwrite != 'y':
                    raise FileExistsError(
                        f"Output file '{output_path}' already exists and overwrite is not allowed.")

        if template_path.endswith((".docx", ".doc")):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for placeholder, value in placeholders.items():
                    paragraph.text = paragraph.text.replace(f"{bookend}{placeholder}{bookend}", value)
            doc.save(output_path)
        elif template_path.endswith(".txt"):
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            for placeholder, value in placeholders.items():
                template_content = template_content.replace(f"{bookend}{placeholder}{bookend}", value)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(template_content)
        else:
            raise ValueError("Unsupported file type. Please use .docx, .doc, or .txt.")

    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {template_path}")
    except FileExistsError as e:
        raise e
    except Exception as e:
        raise IOError(f"An error occurred while generating the document: {e}")


def load_config(config_path):
    """Loads a JSON configuration file."""
    try:
        with open(config_path, 'r') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError, KeyError, IOError) as e:
        print(f"Error loading or processing config file: {e}")
        return None  # Return None to indicate failure
    except Exception as e:
        raise IOError(f"An error occurred while loading the config file: {e}")


def extract_placeholders(template_path, bookends):
    """Extracts placeholders from a template file using specified bookends."""
    placeholders = {}
    try:
        if template_path.endswith((".docx", ".doc")):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                matches = re.findall(rf"{re.escape(bookends)}(.*?){re.escape(bookends)}", paragraph.text)
                for match in matches:
                    placeholder_name = match.strip()
                    if placeholder_name:
                        placeholders[placeholder_name] = ""
        elif template_path.endswith(".txt"):
            with open(template_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
                matches = re.findall(rf"{re.escape(bookends)}(.*?){re.escape(bookends)}", file_content, re.DOTALL)
                for match in matches:
                    placeholder_name = match.strip()
                    if placeholder_name:
                        placeholders[placeholder_name] = ""
        else:
            raise ValueError("Unsupported file type. Please use .docx, .doc, or .txt.")
        return placeholders
    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {template_path}")
    except Exception as e:
        raise IOError(f"An error occurred while extracting placeholders: {e}")


def build_config_from_template(template_path, bookends="%"):
    """Builds a config file from a template file."""
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        filename_parts = os.path.basename(template_path).split('-')
        config_filename = '-'.join(filename_parts[:2]) + "-config.json"
        config_filepath = os.path.join(os.path.dirname(template_path), config_filename)
        date = date.today().strftime("%Y-%m-%d")
        placeholders = extract_placeholders(template_path, bookends)
        config_data = {
            "configFileName": config_filename,
            "templateFilePath": template_path,
            "outputFilePath": config_filename.replace("-config.json", f"-Cover Letter-{date}.docx"),
            "overwriteOutput": False,
            "bookends": bookends,
            "placeholders": placeholders
        }
        placeholders["Date"] = date
        with open(config_filepath, 'w') as outfile:
            json.dump(config_data, outfile, indent=4)
        return config_filepath

    except (FileNotFoundError, ValueError, IOError) as e:
        print(f"Error building config file: {e}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None


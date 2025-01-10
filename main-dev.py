import argparse
import json
import os
from datetime import date
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox


def generate_document(template_path, output_path, data):
    """Generates a document, replacing placeholders and resizing text to fit.

    Supports .docx and .txt files.

    Args:
        template_path: Path to the template file.
        output_path: Path to save the generated document.
         Dictionary containing placeholder names and their replacements.
    """
    try:
        if template_path.endswith(".docx"):
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    paragraph.text = paragraph.text.replace(f"%{key}%", value)
            doc.save(output_path)

        elif template_path.endswith(".txt"):
            with open(template_path, "r") as infile, open(output_path, "w") as outfile:
                for line in infile:
                    for key, value in data.items():
                        line = line.replace(f"%{key}%", value)
                    outfile.write(line)
    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {template_path}")
    except ValueError as e:
        raise
    except Exception as e:
        raise IOError(f"Error processing template: {e}")

    # Check for unmatched placeholders
    if template_path.endswith(".docx"):
        unmatched = [key for paragraph in doc.paragraphs for key in data if f"%{key}%" in paragraph.text]
    elif template_path.endswith(".txt"):
        with open(output_path, "r") as outfile:
            unmatched = [key for line in outfile for key in data if f"%{key}%" in line]
    if unmatched:
        raise ValueError(f"Unmatched placeholders found: {', '.join(unmatched)}")


def load_config(path):
    """Loads config from a JSON file. Returns an empty dictionary if the file doesn't exist."""
    try:
        with open(path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


def browse_config_file():
    """Open a file dialog to select the config.json file."""
    file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
    if file_path:
        global config_path  # Declare config_path_var as global
        config_path.set(file_path)
        load_config_file()


def browse_template_file():
    """Open a file dialog to select the template file."""
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("Text files", "*.txt")])
    if file_path:
        template_path.set(file_path)


def load_config_file():
    """Load the content of the selected config.json file and populate the corresponding fields."""
    global config_path_var  # Declare config_path_var as global
    config_path = config_path_var.get()
    if not os.path.exists(config_path):
        messagebox.showerror("Error", "Config file not found.")
        return
    config = load_config(config_path)
    # Populate fields with config data (this is a placeholder, adjust as needed)
    template_path.set(config.get("templateFilePath", ""))
    output_filename.set(config.get("outputFilePath", ""))
    placeholders.set(json.dumps(config.get("placeholders", {}), indent=4))


def add_placeholder():
    global placeholder_var, replace_with_var, placeholders_text
    placeholder = placeholder_var.get()
    replace_with = replace_with_var.get()
    if placeholder and replace_with:
        placeholders_text.insert(tk.END, f"{placeholder}: {replace_with}\n")
        placeholder_var.set("")
        replace_with_var.set("")
    else:
        messagebox.showerror("Error", "Both fields must be filled out.")


def run_gui():
    global config_path  # Declare config_path_var as global
    global template_path
    global output_filename
    global placeholders
    # Create the main window
    root = tk.Tk()
    root.title("Job Application Helper")

    labelCol = 0
    buttonA = 1
    inputCol = 2
    buttonB = 4

    # Config file path
    tk.Label(root, text="Config file path:").grid(row=0, column=labelCol,
                                                  padx=5, pady=5, sticky="e")
    tk.Button(root, text="Browse", command=browse_config_file).grid(row=0, column=buttonA,
                                                                    padx=5, pady=5, sticky="w")
    config_path = tk.StringVar()
    tk.Entry(root, textvariable=config_path, width=50).grid(row=0, column=inputCol,
                                                                padx=5, pady=5, sticky="w")
    tk.Button(root, text="Load Config", command=load_config_file).grid(row=0, column=buttonB,
                                                                       padx=5, pady=5, sticky="w")

    # Template file path
    tk.Label(root, text="Template file path:").grid(row=1, column=labelCol,
                                                    padx=5, pady=5, sticky="e")
    tk.Button(root, text="Browse", command=browse_template_file).grid(row=1, column=buttonA,
                                                                      padx=5, pady=5, sticky="w")
    template_path = tk.StringVar()
    tk.Entry(root, textvariable=template_path, width=50).grid(row=1, column=inputCol,
                                                              padx=5, pady=5, sticky="w")

    # Output file path
    tk.Label(root, text="Output file path:").grid(row=2, column=labelCol,
                                                  padx=5, pady=5, sticky="e")
    output_filename = tk.StringVar()
    tk.Entry(root, textvariable=output_filename, width=50).grid(row=2, column=inputCol, columnspan=2,
                                                                padx=5, pady=5, sticky="w")

    # Placeholder key: value pair boxes
    tk.Label(root, text="Placeholder (e.g. %FirstName%):").grid(row=3, column=labelCol,
                                                                padx=5, pady=5, sticky="e")
    placeholder_var = tk.StringVar()
    tk.Entry(root, textvariable=placeholder_var, width=50).grid(row=3, column=inputCol, columnspan=1,
                                                                padx=5, pady=5, sticky="w")
    tk.Label(root, text="Replace with:").grid(row=4, column=labelCol,
                                              padx=5, pady=5, sticky="e")
    replace_with_var = tk.StringVar()
    tk.Entry(root, textvariable=replace_with_var, width=50).grid(row=4, column=inputCol, columnspan=1,
                                                                 padx=5, pady=5, sticky="w")
    tk.Button(root, text="Add", command=add_placeholder).grid(row=4, column=buttonB, padx=5, pady=5, sticky="w")

    # Placeholders
    tk.Label(root, text="Placeholders:").grid(row=5, column=labelCol,
                                              padx=5, pady=5, sticky="e")
    placeholders_text = tk.Text(root, width=50, height=3, wrap="word")
    placeholders_text.grid(row=5, rowspan=3,
                           column=inputCol, columnspan=3,
                           padx=5, pady=5, sticky="w")

    # Scrollbar for the placeholders text box
    scrollbar = tk.Scrollbar(root, command=placeholders_text.yview)
    scrollbar.grid(row=5, rowspan=3, column=inputCol + 1, sticky="nwes")
    placeholders_text.config(yscrollcommand=scrollbar.set)

    # Adjust column weights to minimize whitespace
    root.grid_columnconfigure(0, weight=0)
    root.grid_columnconfigure(1, weight=0)
    root.grid_columnconfigure(2, weight=1)
    root.grid_columnconfigure(3, weight=0)
    root.grid_columnconfigure(4, weight=0)


    # Run the main loop
    root.mainloop()


if __name__ == "__main__":
    placeholder_var = ""
    replace_with_var = ""

    # Parse command-line arguments
    parser = argparse.ArgumentParser(description="Generate documents from templates.")
    parser.add_argument("--config", help="Path to the config file (default: config.json)")
    parser.add_argument("-GUI", action="store_true", help="Load the GUI")
    args = parser.parse_args()

    if args.GUI:
        run_gui()
    else:
        config_path = args.config if args.config else input(
            "Enter the path to the config file (default: config.json): ") or "config.json"
        print(f"Loading config from {config_path}...")
        config = load_config(config_path)

        # Automatically generate date
        today = date.today().strftime("%Y-%m-%d")

        while True:
            template_path = config.get("templateFilePath")
            if template_path is None:
                template_path = input("Enter template file path: ")
            if not os.path.exists(template_path):
                print("Error: Template file not found. Please enter a valid path.")
                if template_path is not None:  # Only clear the config value if it was from the file
                    config["template_path"] = None  # Clear incorrect value from config
                continue  # Ask for input again
            print(f"Using template file: {template_path}")
            break

        while True:
            output_filename = config.get("outputFilePath")
            if output_filename is None:
                output_filename = input("Enter output file name: ")

            overwrite = config.get("overwriteOutput")  # Initialize overwrite variable
            if overwrite is None:
                overwrite = False

            while os.path.exists(output_filename):
                overwrite = input(f"File '{output_filename}' already exists. Overwrite? (y/n): ")
                if overwrite:
                    break  # Exit the loop if user wants to overwrite
                elif not overwrite:
                    output_filename = input("Enter a different output file name: ")  # Ask for a new name
                    overwrite = None  # Reset overwrite for the new filename check
                else:
                    print("Invalid input. Please enter 'y' or 'n'.")  # Handle invalid input

            if overwrite and output_filename is not None:
                break  # Exit the loop if user wants to overwrite
            else:
                continue  # Ask for input again if not overwriting

        print(f"Output file will be saved as: {output_filename}")

        placeholders = config.get("placeholders", {})  # Load placeholders from config, default to empty dict

        if not placeholders:  # Only ask for placeholders if data is empty
            print("No data found in config. Please enter placeholder values.")
            while True:
                key = input("Enter placeholder name (or type 'done'): ")
                if key == "done":
                    break
                value = input(f"Enter value for {key}: ")
                placeholders[key] = value

        # Add the automatically generated date to the data dictionary
        placeholders["date"] = today

        try:
            print("Generating document...")
            generate_document(template_path, output_filename, placeholders)
            print(f"Document '{output_filename}' generated successfully.")
        except Exception as e:
            print(f"Error: {e}")
